from __future__ import annotations

import json
import sqlite3
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Callable, Iterable

from docx import Document
from PIL import Image, ImageEnhance, ImageOps
from pypdf import PdfReader

from core.documents import DocumentManager
from core.projects import ProjectInfo

from .normalization import normalize_extracted_text, text_blocks
from .quality import TextQuality, evaluate_text


def _utc_now() -> str:
    return datetime.now(timezone.utc).isoformat()


PAGE_STATES = (
    "native_text_accepted", "native_text_low_quality", "ocr_queued", "ocr_completed",
    "ocr_retry_completed", "partial_text", "unreadable_after_retry", "user_review_required",
)


@dataclass(frozen=True, slots=True)
class OCRResult:
    document_id: str
    text_path: Path
    metadata_path: Path
    page_count: int
    character_count: int
    method: str
    completed_at: str


@dataclass(frozen=True, slots=True)
class OCRPageResult:
    document_id: str
    page_number: int
    raw_text: str
    normalized_text: str
    corrected_text: str
    extraction_method: str
    state: str
    confidence: float
    quality: dict[str, object]
    blocks: list[dict[str, object]]
    failure_reason: str
    rotation: int
    user_readability: str
    updated_at: str

    @property
    def downstream_text(self) -> str:
        return self.corrected_text or self.normalized_text or self.raw_text


class OCRManager:
    """Page-aware native-first extraction with bounded local OCR recovery."""

    def __init__(self, project: ProjectInfo) -> None:
        self.project = project
        self.documents = DocumentManager(project)
        self.output_dir = project.root / "ocr"
        self.output_dir.mkdir(parents=True, exist_ok=True)
        self._ensure_schema()

    def process(
        self,
        document_id: str,
        *,
        progress: Callable[[int, str], None] | None = None,
        cancelled: Callable[[], bool] | None = None,
        force_ocr: bool = False,
    ) -> OCRResult:
        document = self.documents.get(document_id)
        emit = progress or (lambda _percent, _message: None)
        is_cancelled = cancelled or (lambda: False)
        emit(1, f"Preparing {document.original_name}")
        suffix = document.stored_path.suffix.lower()
        acroform = {}

        if suffix == ".pdf":
            pages, acroform = self._extract_pdf(
                document.stored_path, emit, is_cancelled, force_ocr=force_ocr
            )
        elif suffix in {".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp"}:
            if is_cancelled():
                raise InterruptedError("OCR cancelled")
            with Image.open(document.stored_path) as image:
                pages = [self._ocr_image_page(image.copy(), 1)]
        elif suffix == ".docx":
            word = Document(document.stored_path)
            raw = "\n".join(paragraph.text for paragraph in word.paragraphs)
            pages = [self._text_page(1, raw, "docx-text")]
        elif suffix in {".txt", ".md", ".csv", ".json"}:
            raw = document.stored_path.read_text(encoding="utf-8", errors="replace")
            pages = [self._text_page(1, raw, "plain-text")]
        else:
            raise ValueError(f"OCR/text extraction is not supported for {suffix or 'this file type'}.")
        if is_cancelled():
            raise InterruptedError("OCR cancelled")

        self._persist_pages(document_id, pages)
        # Reanalysis preserves user corrections in the database; downstream output must use them.
        pages = self.list_pages(document_id)
        combined = "\n".join(
            f"--- PAGE {page.page_number} ---\n{page.downstream_text.strip()}\n"
            for page in pages
        ).strip() + "\n"
        text_path = self.output_dir / f"{document_id}.txt"
        metadata_path = self.output_dir / f"{document_id}.json"
        text_path.write_text(combined, encoding="utf-8")
        methods = sorted({page.extraction_method for page in pages})
        completed_at = _utc_now()
        metadata = {
            "document_id": document_id,
            "source_name": document.original_name,
            "page_count": len(pages),
            "character_count": len(combined),
            "method": "+".join(methods),
            "completed_at": completed_at,
            "text_path": str(text_path),
            "acroform_fields": acroform,
            "pages": [
                {
                    "page_number": page.page_number, "state": page.state,
                    "method": page.extraction_method, "confidence": page.confidence,
                    "failure_reason": page.failure_reason,
                }
                for page in pages
            ],
        }
        metadata_path.write_text(json.dumps(metadata, indent=2), encoding="utf-8")
        with self._connect() as connection:
            connection.execute(
                """INSERT OR REPLACE INTO ocr_results(
                    document_id,text_path,metadata_path,page_count,character_count,method,completed_at
                ) VALUES(?,?,?,?,?,?,?)""",
                (
                    document_id, str(text_path.relative_to(self.project.root)),
                    str(metadata_path.relative_to(self.project.root)), len(pages),
                    len(combined), "+".join(methods), completed_at,
                ),
            )
            connection.execute(
                "UPDATE documents SET status=? WHERE document_id=?",
                ("ocr_complete" if any(page.downstream_text.strip() for page in pages) else "ocr_failed", document_id),
            )
        emit(100, f"Completed {document.original_name}")
        return self.get(document_id)

    def get(self, document_id: str) -> OCRResult:
        with self._connect() as connection:
            row = connection.execute(
                """SELECT document_id,text_path,metadata_path,page_count,
                   character_count,method,completed_at FROM ocr_results WHERE document_id=?""",
                (document_id,),
            ).fetchone()
        if row is None:
            raise KeyError(document_id)
        return OCRResult(
            str(row[0]), self.project.root / str(row[1]), self.project.root / str(row[2]),
            int(row[3]), int(row[4]), str(row[5]), str(row[6]),
        )

    def list_results(self) -> list[OCRResult]:
        with self._connect() as connection:
            ids = [row[0] for row in connection.execute(
                "SELECT document_id FROM ocr_results ORDER BY completed_at DESC,document_id"
            )]
        return [self.get(str(document_id)) for document_id in ids]

    def list_pages(self, document_id: str) -> list[OCRPageResult]:
        with self._connect() as connection:
            rows = connection.execute(
                "SELECT * FROM ocr_pages WHERE document_id=? ORDER BY page_number",
                (document_id,),
            ).fetchall()
        return [self._page_from_row(row) for row in rows]

    def get_page(self, document_id: str, page_number: int) -> OCRPageResult:
        with self._connect() as connection:
            row = connection.execute(
                "SELECT * FROM ocr_pages WHERE document_id=? AND page_number=?",
                (document_id, page_number),
            ).fetchone()
        if row is None:
            raise KeyError((document_id, page_number))
        return self._page_from_row(row)

    def save_correction(
        self, document_id: str, page_number: int, text: str, *, readability: str = "readable"
    ) -> OCRPageResult:
        if readability not in {"", "readable", "unreadable"}:
            raise ValueError("Unsupported readability override")
        with self._connect() as connection:
            if not connection.execute(
                """UPDATE ocr_pages SET corrected_text=?,user_readability=?,state=?,
                   updated_at=? WHERE document_id=? AND page_number=?""",
                (
                    text, readability,
                    "user_review_required" if readability == "" else (
                        "unreadable_after_retry" if readability == "unreadable" else "ocr_completed"
                    ),
                    _utc_now(), document_id, page_number,
                ),
            ).rowcount:
                raise KeyError((document_id, page_number))
        self._rewrite_combined(document_id)
        return self.get_page(document_id, page_number)

    def revert_correction(self, document_id: str, page_number: int) -> OCRPageResult:
        with self._connect() as connection:
            if not connection.execute(
                "UPDATE ocr_pages SET corrected_text='',user_readability='',updated_at=? "
                "WHERE document_id=? AND page_number=?",
                (_utc_now(), document_id, page_number),
            ).rowcount:
                raise KeyError((document_id, page_number))
        self._rewrite_combined(document_id)
        return self.get_page(document_id, page_number)

    def retry_page(self, document_id: str, page_number: int, *, rotation: int = 0) -> OCRPageResult:
        document = self.documents.get(document_id)
        if document.stored_path.suffix.lower() == ".pdf":
            image = self._render_pdf_page(document.stored_path, page_number - 1, 300)
        else:
            image = Image.open(document.stored_path).copy()
        if rotation % 360:
            image = image.rotate(-(rotation % 360), expand=True)
        page = self._ocr_image_page(image, page_number, retry=True, rotation=rotation % 360)
        self._persist_pages(document_id, [page])
        self._rewrite_combined(document_id)
        return self.get_page(document_id, page_number)

    def delete_outputs(self, document_ids: Iterable[str]) -> int:
        ids = list(dict.fromkeys(str(item) for item in document_ids))
        if not ids:
            return 0
        placeholders = ",".join("?" for _ in ids)
        with self._connect() as connection:
            existing = connection.execute(
                f"SELECT document_id,text_path,metadata_path FROM ocr_results WHERE document_id IN ({placeholders})",
                ids,
            ).fetchall()
            connection.execute(f"DELETE FROM ocr_results WHERE document_id IN ({placeholders})", ids)
            connection.execute(f"DELETE FROM ocr_pages WHERE document_id IN ({placeholders})", ids)
            connection.execute(
                f"UPDATE documents SET status='imported' WHERE document_id IN ({placeholders})", ids
            )
        for _document_id, text_path, metadata_path in existing:
            for relative in (text_path, metadata_path):
                path = self.project.root / str(relative)
                path.unlink(missing_ok=True)
        return len(existing)

    def clear_failed_jobs(self, document_ids: Iterable[str]) -> int:
        ids = list(dict.fromkeys(str(item) for item in document_ids))
        if not ids:
            return 0
        with self._connect() as connection:
            count = 0
            for document_id in ids:
                count += connection.execute(
                    "DELETE FROM jobs WHERE job_type='ocr' AND status='failed' "
                    "AND json_extract(payload_json,'$.document_id')=?",
                    (document_id,),
                ).rowcount
        return count

    def _extract_pdf(self, path, emit, is_cancelled, *, force_ocr=False):
        reader = PdfReader(str(path))
        acroform = self._acroform_fields(reader)
        pages: list[OCRPageResult] = []
        total = max(len(reader.pages), 1)
        for index, source_page in enumerate(reader.pages, start=1):
            if is_cancelled():
                raise InterruptedError("OCR cancelled")
            emit(max(2, int(index / total * 90)), f"Extracting PDF page {index} of {total}")
            try:
                raw = source_page.extract_text(extraction_mode="layout") or source_page.extract_text() or ""
            except Exception:
                raw = ""
            native = self._text_page(index, raw, "pypdf-layout")
            if not force_ocr and self._native_usable(native.quality):
                pages.append(native)
                continue
            try:
                image = self._render_pdf_page(path, index - 1, 300)
                ocr_page = self._ocr_image_page(image, index)
                pages.append(ocr_page if ocr_page.confidence > native.confidence else self._low_native(native))
            except Exception as exc:
                reason = self._safe_ocr_error(exc)
                pages.append(self._low_native(native, failure_reason=reason))
        return pages, acroform

    @staticmethod
    def _acroform_fields(reader: PdfReader) -> dict[str, object]:
        try:
            fields = reader.get_fields() or {}
        except Exception:
            return {}
        output: dict[str, object] = {}
        for name, field in fields.items():
            value = field.get("/V")
            if value not in (None, "", "/Off"):
                output[str(name)] = str(value).lstrip("/")
        return output

    @staticmethod
    def _native_usable(quality: dict[str, object]) -> bool:
        return (
            int(quality.get("character_count", 0)) >= 30
            and int(quality.get("meaningful_words", 0)) >= 5
            and float(quality.get("alphanumeric_ratio", 0)) >= 0.55
            and float(quality.get("score", 0)) >= 0.48
        )

    def _text_page(self, number: int, raw: str, method: str) -> OCRPageResult:
        normalized = normalize_extracted_text(raw)
        quality = evaluate_text(normalized)
        state = "native_text_accepted" if self._native_usable(quality.to_dict()) else (
            "partial_text" if normalized.strip() else "native_text_low_quality"
        )
        return OCRPageResult(
            "", number, raw, normalized, "", method, state, quality.score,
            quality.to_dict(), text_blocks(raw), "", 0, "", _utc_now(),
        )

    def _low_native(self, page: OCRPageResult, failure_reason: str = "") -> OCRPageResult:
        state = "partial_text" if page.raw_text.strip() else "unreadable_after_retry"
        return OCRPageResult(
            "", page.page_number, page.raw_text, page.normalized_text, "",
            page.extraction_method, state, page.confidence, page.quality, page.blocks,
            failure_reason or "Native text was low quality and OCR did not produce a better result.",
            0, "", _utc_now(),
        )

    def _ocr_image_page(
        self, image: Image.Image, page_number: int, *, retry: bool = False, rotation: int = 0
    ) -> OCRPageResult:
        import pytesseract

        candidates: list[tuple[str, str, float]] = []
        profiles = (
            ("grayscale-autocontrast", self._preprocess(image, contrast=1.25), 6),
            ("grayscale-sparse", self._preprocess(image, contrast=1.05), 11),
        )
        errors: list[str] = []
        for profile, prepared, psm in profiles:
            try:
                data = pytesseract.image_to_data(
                    prepared, config=f"--oem 3 --psm {psm}",
                    output_type=pytesseract.Output.DICT,
                )
                words = [str(value) for value in data.get("text", [])]
                raw = self._reconstruct_tesseract(data, words)
                confidences = [
                    float(value) for value in data.get("conf", [])
                    if str(value).replace(".", "", 1).lstrip("-").isdigit() and float(value) >= 0
                ]
                engine_conf = sum(confidences) / len(confidences) if confidences else -1.0
                candidates.append((profile, raw, engine_conf))
            except Exception as exc:
                errors.append(self._safe_ocr_error(exc))
        if not candidates:
            quality = evaluate_text("")
            return OCRPageResult(
                "", page_number, "", "", "", "tesseract", "unreadable_after_retry", 0,
                quality.to_dict(), [], "; ".join(dict.fromkeys(errors)), rotation, "", _utc_now(),
            )
        ranked: list[tuple[float, str, str, float, TextQuality]] = []
        for profile, raw, engine_conf in candidates:
            normalized = normalize_extracted_text(raw)
            quality = evaluate_text(normalized, engine_confidence=engine_conf)
            ranked.append((quality.score, profile, raw, engine_conf, quality))
        _score, profile, raw, engine_conf, quality = max(ranked, key=lambda item: item[0])
        normalized = normalize_extracted_text(raw)
        state = "ocr_retry_completed" if retry else "ocr_completed"
        if quality.score < 0.48:
            state = "partial_text" if normalized else "unreadable_after_retry"
        elif quality.review_required:
            state = "user_review_required"
        return OCRPageResult(
            "", page_number, raw, normalized, "", f"tesseract-{profile}",
            state, quality.score, quality.to_dict(), text_blocks(raw),
            "", rotation, "", _utc_now(),
        )

    @staticmethod
    def _preprocess(image: Image.Image, *, contrast: float) -> Image.Image:
        grayscale = ImageOps.grayscale(image)
        grayscale = ImageOps.autocontrast(grayscale, cutoff=1)
        return ImageEnhance.Contrast(grayscale).enhance(contrast)

    @staticmethod
    def _reconstruct_tesseract(data: dict, words: list[str]) -> str:
        lines: list[str] = []
        current_key = None
        current: list[str] = []
        for index, word in enumerate(words):
            word = word.strip()
            key = tuple(data.get(name, [0] * len(words))[index] for name in ("block_num", "par_num", "line_num"))
            if current_key is not None and key != current_key:
                if current:
                    lines.append(" ".join(current))
                current = []
            current_key = key
            if word:
                current.append(word)
        if current:
            lines.append(" ".join(current))
        return "\n".join(lines)

    @staticmethod
    def _render_pdf_page(path: Path, page_index: int, dpi: int) -> Image.Image:
        try:
            import pypdfium2 as pdfium
        except ImportError as exc:
            raise RuntimeError(
                "Scanned PDF OCR requires the bundled pypdfium2 renderer."
            ) from exc
        pdf = pdfium.PdfDocument(str(path))
        try:
            bitmap = pdf[page_index].render(scale=dpi / 72.0)
            return bitmap.to_pil().copy()
        finally:
            pdf.close()

    def _persist_pages(self, document_id: str, pages: list[OCRPageResult]) -> None:
        with self._connect() as connection:
            for page in pages:
                connection.execute(
                    """INSERT INTO ocr_pages(
                       document_id,page_number,raw_text,normalized_text,corrected_text,
                       extraction_method,state,confidence,quality_json,blocks_json,
                       failure_reason,rotation,user_readability,created_at,updated_at
                    ) VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)
                    ON CONFLICT(document_id,page_number) DO UPDATE SET
                       raw_text=excluded.raw_text,normalized_text=excluded.normalized_text,
                       extraction_method=excluded.extraction_method,state=CASE
                         WHEN ocr_pages.user_readability<>'' THEN ocr_pages.state ELSE excluded.state END,
                       confidence=excluded.confidence,quality_json=excluded.quality_json,
                       blocks_json=excluded.blocks_json,failure_reason=excluded.failure_reason,
                       rotation=excluded.rotation,updated_at=excluded.updated_at""",
                    (
                        document_id, page.page_number, page.raw_text, page.normalized_text,
                        page.corrected_text, page.extraction_method, page.state, page.confidence,
                        json.dumps(page.quality), json.dumps(page.blocks), page.failure_reason,
                        page.rotation, page.user_readability, _utc_now(), _utc_now(),
                    ),
                )

    def _rewrite_combined(self, document_id: str) -> None:
        result = self.get(document_id)
        pages = self.list_pages(document_id)
        combined = "\n".join(
            f"--- PAGE {page.page_number} ---\n{page.downstream_text.strip()}\n" for page in pages
        ).strip() + "\n"
        result.text_path.write_text(combined, encoding="utf-8")
        with self._connect() as connection:
            connection.execute(
                "UPDATE ocr_results SET character_count=?,completed_at=? WHERE document_id=?",
                (len(combined), _utc_now(), document_id),
            )

    def _connect(self) -> sqlite3.Connection:
        connection = sqlite3.connect(self.project.database_path)
        connection.row_factory = sqlite3.Row
        connection.execute("PRAGMA foreign_keys=ON")
        return connection

    @staticmethod
    def _page_from_row(row: sqlite3.Row) -> OCRPageResult:
        return OCRPageResult(
            row["document_id"], int(row["page_number"]), row["raw_text"],
            row["normalized_text"], row["corrected_text"], row["extraction_method"],
            row["state"], float(row["confidence"]), json.loads(row["quality_json"]),
            json.loads(row["blocks_json"]), row["failure_reason"], int(row["rotation"]),
            row["user_readability"], row["updated_at"],
        )

    @staticmethod
    def _safe_ocr_error(exc: Exception) -> str:
        message = str(exc)
        if type(exc).__name__ == "TesseractNotFoundError":
            return "Tesseract OCR engine was not found. Install Tesseract or use native-text documents."
        return (message or type(exc).__name__)[:500]

    def _ensure_schema(self) -> None:
        with self._connect() as connection:
            connection.executescript(
                """
                CREATE TABLE IF NOT EXISTS ocr_results (
                    document_id TEXT PRIMARY KEY,text_path TEXT NOT NULL,metadata_path TEXT NOT NULL,
                    page_count INTEGER NOT NULL,character_count INTEGER NOT NULL,method TEXT NOT NULL,
                    completed_at TEXT NOT NULL,
                    FOREIGN KEY(document_id) REFERENCES documents(document_id) ON DELETE CASCADE
                );
                CREATE TABLE IF NOT EXISTS ocr_pages (
                    document_id TEXT NOT NULL,page_number INTEGER NOT NULL,raw_text TEXT NOT NULL DEFAULT '',
                    normalized_text TEXT NOT NULL DEFAULT '',corrected_text TEXT NOT NULL DEFAULT '',
                    extraction_method TEXT NOT NULL,state TEXT NOT NULL,confidence REAL NOT NULL DEFAULT 0,
                    quality_json TEXT NOT NULL DEFAULT '{}',blocks_json TEXT NOT NULL DEFAULT '[]',
                    failure_reason TEXT NOT NULL DEFAULT '',rotation INTEGER NOT NULL DEFAULT 0,
                    user_readability TEXT NOT NULL DEFAULT '',created_at TEXT NOT NULL,updated_at TEXT NOT NULL,
                    PRIMARY KEY(document_id,page_number),
                    FOREIGN KEY(document_id) REFERENCES documents(document_id) ON DELETE CASCADE
                );
                CREATE INDEX IF NOT EXISTS idx_ocr_completed_at ON ocr_results(completed_at);
                CREATE INDEX IF NOT EXISTS idx_ocr_page_review ON ocr_pages(state,confidence);
                """
            )
