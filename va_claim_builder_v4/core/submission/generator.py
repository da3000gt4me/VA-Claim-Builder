from __future__ import annotations
import csv,hashlib,json,shutil,tempfile,uuid,zipfile
from datetime import datetime,timezone
from pathlib import Path
from core.claims import ClaimManager
from core.dbq import DBQManager
from core.documents import DocumentManager
from core.evidence import EvidenceManager
from core.nexus import NexusLetterManager
from core.optimizer import OptimizerManager
from core.rating_strategy import RatingStrategyManager
from core.timeline import TimelineManager
from .manager import SubmissionManager
def _now():return datetime.now(timezone.utc).isoformat()
class SubmissionGenerationError(RuntimeError):pass
class SubmissionGenerator:
 def __init__(self,project):self.project=project;self.manager=SubmissionManager(project);self.claims=ClaimManager(project);self.documents=DocumentManager(project);self.evidence=EvidenceManager(project);self.timeline=TimelineManager(project);self.nexus=NexusLetterManager(project);self.dbq=DBQManager(project);self.rating=RatingStrategyManager(project);self.optimizer=OptimizerManager(project)
 def generate(self,pid,*,job_id=None,allow_incomplete=False,cancelled=lambda:False,progress=lambda p,m:None):
  p=self.manager.get(pid);issues=self.manager.validate(pid,self.project.root/"reports"/"submissions");blocking=[x for x in issues if x.level=="blocking"]
  if blocking and not allow_incomplete:raise SubmissionGenerationError("Blocking validation errors must be resolved before generation")
  version=p.package_version+1;export_id=str(uuid.uuid4());created=_now()
  with self.manager._connect() as c:c.execute("UPDATE submission_packages SET status='generating',package_version=?,updated_at=? WHERE package_id=?",(version,created,pid));c.execute("INSERT INTO submission_exports VALUES(?,?,?,'pending','',?,?,'{}','{}',?,'',?,NULL)",(export_id,pid,version,json.dumps([]),json.dumps([x.__dict__ if hasattr(x,'__dict__') else {k:getattr(x,k) for k in x.__slots__} for x in issues]),job_id,created))
  temp=Path(tempfile.mkdtemp(prefix="submission-",dir=self.project.root/"temp"));final=None
  try:
   if cancelled():raise InterruptedError("Package generation cancelled")
   progress(10,"Building summary binder");p=self.manager.get(pid);binder=temp/"submission-summary.docx";self._binder(p,binder,issues);index_docx=temp/"evidence-index.docx";self._index_docx(p,index_docx);index_csv=temp/"evidence-index.csv";self._index_csv(p,index_csv)
   companion=temp/"companion-sources";companion.mkdir();pdfs=[];source_files=[];unmerged=[]
   progress(35,"Collecting selected source files")
   for s in p.sources:
    if s.source_type not in {"document","claimant_statement","witness_statement","provider_request"}:continue
    d=self.documents.get(s.source_id)
    if not d.stored_path.is_file():continue
    target=companion/f"{s.exhibit_id}-{self._safe(d.original_name)}";shutil.copy2(d.stored_path,target);source_files.append((s,target,d.sha256))
    if target.suffix.lower()==".pdf":pdfs.append((s,target))
    else:unmerged.append(target.name)
   if cancelled():raise InterruptedError("Package generation cancelled")
   progress(55,"Merging readable PDF exhibits");merged=temp/"consolidated-evidence.pdf";merge_fail=[]
   if pdfs:
    from pypdf import PdfReader,PdfWriter
    writer=PdfWriter()
    for s,path in pdfs:
     try:
      reader=PdfReader(path,strict=False)
      if reader.is_encrypted and reader.decrypt("")==0:raise ValueError("encrypted")
      for page in reader.pages:writer.add_page(page)
     except Exception:merge_fail.append(path.name)
    if len(writer.pages):
     with merged.open("wb") as f:writer.write(f)
    else:merged=None
   else:merged=None
   unmerged+=merge_fail;progress(75,"Writing package manifest")
   outputs=[binder.name,index_docx.name,index_csv.name]+([merged.name] if merged else [])+["companion-sources/"+x[1].name for x in source_files]
   manifest={"package_name":p.name,"package_type":p.package_type,"package_version":version,"generation_timestamp":_now(),"included_claims":[self.claims.get(cid).condition_name for cid in p.claim_ids],"sections":[{"key":s["section_key"],"title":s["title"],"order":s["sort_order"]} for s in p.sections if s["enabled"]],"exhibits":[{"exhibit":s.exhibit_id,"type":s.source_type,"title":s.display_name or self._source_name(s),"section":s.section_key} for s in p.sources],"output_filenames":outputs,"source_checksums":[{"exhibit":s.exhibit_id,"sha256":checksum} for s,_,checksum in source_files],"validation_summary":{"blocking":sum(i.level=="blocking" for i in issues),"warnings":sum(i.level=="warning" for i in issues),"information":sum(i.level=="info" for i in issues)},"unmerged_files":unmerged,"advisory":"This locally generated organization aid is not legal representation and does not guarantee VA approval."};(temp/"package-manifest.json").write_text(json.dumps(manifest,indent=2),encoding="utf-8");outputs.append("package-manifest.json")
   if cancelled():raise InterruptedError("Package generation cancelled")
   base=self.project.root/"reports"/"submissions";base.mkdir(parents=True,exist_ok=True);final=base/f"{self._safe(p.name)}-v{version}"
   if final.exists():final=base/f"{self._safe(p.name)}-v{version}-{datetime.now():%Y%m%d%H%M%S}"
   temp.replace(final);zip_path=final.with_suffix(".zip")
   with zipfile.ZipFile(zip_path,"w",zipfile.ZIP_DEFLATED) as z:
    for f in final.rglob("*"):
     if f.is_file():z.write(f,f.relative_to(final.parent))
   progress(100,"Package generation complete");completed=_now();formats=["docx_binder","docx_index","csv_index","folder","zip","json_manifest"]+(["consolidated_pdf"] if merged else [])
   with self.manager._connect() as c:c.execute("UPDATE submission_packages SET status=?,generated_at=?,updated_at=? WHERE package_id=?",("incomplete" if blocking else "completed",completed,completed,pid));c.execute("UPDATE submission_exports SET status='completed',output_path=?,formats_json=?,manifest_json=?,generation_metadata_json=?,completed_at=? WHERE export_id=?",(str(final),json.dumps(formats),json.dumps(manifest),json.dumps({"local_only":True,"atomic_finalize":True}),completed,export_id))
   self.manager._snapshot(pid);return {"export_id":export_id,"output_folder":final,"zip_path":zip_path,"manifest":manifest,"issues":issues}
  except InterruptedError as exc:
   shutil.rmtree(temp,ignore_errors=True);self._failed(pid,export_id,"cancelled",str(exc));return {"export_id":export_id,"status":"cancelled"}
  except Exception as exc:
   shutil.rmtree(temp,ignore_errors=True);self._failed(pid,export_id,"failed",f"Package generation failed: {type(exc).__name__}");raise SubmissionGenerationError(f"Package generation failed: {type(exc).__name__}") from exc
 def _binder(self,p,path,issues):
  from docx import Document
  d=Document();d.add_heading(p.name,0);d.add_paragraph("Locally generated submission organization aid. This package is not legal representation and does not guarantee VA approval or a disability rating.");d.add_paragraph(f"Package type: {p.package_type.replace('_',' ').title()} | Version: {p.package_version}")
  for sec in p.sections:
   if not sec["enabled"]:continue
   d.add_heading(sec["title"],1);key=sec["section_key"]
   if key=="claim_list":
    for cid in p.claim_ids:d.add_paragraph(self.claims.get(cid).condition_name,style="List Bullet")
   elif key=="medical_chronology":
    for cid in p.claim_ids:
     for t in self.timeline.list(claim_id=cid):d.add_paragraph(f"{t.event_date} — {t.title}: {t.description}")
   elif key=="claim_evidence_summaries":self._claim_summaries(d,p)
   elif key in {"evidence_index","exhibit_list"}:
    for s in p.sources:d.add_paragraph(f"{s.exhibit_id} — {s.display_name or self._source_name(s)} ({s.source_type})")
   elif key=="rating_strategy":
    for s in p.sources:
     if s.source_type=="rating_strategy":r=self.rating.get(s.source_id);d.add_paragraph(f"Estimated range: {r.estimated_rating_range}. Advisory; not guaranteed.\n"+r.generated_reasoning)
   elif key=="optimizer_summary":
    for s in p.sources:
     if s.source_type=="optimizer":o=self.optimizer.get(s.source_id);d.add_paragraph(f"Readiness: {o.overall_score}%. "+o.score_explanation.get("disclaimer",""))
   elif key=="contradictory_evidence":
    for s in p.sources:
     if s.source_type=="rating_strategy":
      for item in self.rating.get(s.source_id).contradictory_evidence:d.add_paragraph(item,style="List Bullet")
   elif key=="validation_report":
    for i in issues:d.add_paragraph(f"{i.level.upper()}: {i.message}")
   else:d.add_paragraph("See indexed exhibits and source traceability in this package.")
  d.save(path)
 def _claim_summaries(self,d,p):
  selected_evidence={s.source_id for s in p.sources if s.source_type=="evidence"};selected_timeline={s.source_id for s in p.sources if s.source_type=="timeline"}
  for cid in p.claim_ids:
   c=self.claims.get(cid);d.add_heading(c.condition_name,2);d.add_paragraph(f"Claim type/theory: {c.claim_type}. Condition summary: {c.description or 'Not provided.'}");d.add_paragraph(f"Diagnosis evidence: {c.current_diagnosis or 'Missing or intentionally not included.'}");d.add_paragraph(f"In-service event/exposure/aggravation: {c.service_event or 'Missing or intentionally not included.'}");d.add_paragraph(f"Reported symptoms and functional impact: {c.symptoms or 'Not provided.'}")
   for e in self.evidence.list(claim_id=cid):
    if e.evidence_id in selected_evidence:d.add_paragraph(f"Evidence — {e.title}: {e.description} [review: {e.review_status}]")
   for t in self.timeline.list(claim_id=cid):
    if t.event_id in selected_timeline:d.add_paragraph(f"Timeline — {t.event_date}: {t.title}; {t.functional_impact}")
 def _index_docx(self,p,path):
  from docx import Document
  d=Document();d.add_heading("Evidence Index and Exhibit List",0);table=d.add_table(rows=1,cols=5);hdr=table.rows[0].cells
  for i,x in enumerate(("Exhibit","Title","Type","Section","Traceability")):hdr[i].text=x
  for s in p.sources:
   cells=table.add_row().cells
   for i,x in enumerate((s.exhibit_id,s.display_name or self._source_name(s),s.source_type,s.section_key,"Original project source preserved")):cells[i].text=x
  d.save(path)
 def _index_csv(self,p,path):
  with path.open("w",newline="",encoding="utf-8") as f:
   w=csv.writer(f);w.writerow(("exhibit","title","source_type","section","user_confirmed"));w.writerows((s.exhibit_id,s.display_name or self._source_name(s),s.source_type,s.section_key,s.user_confirmed) for s in p.sources)
 def _source_name(self,s):
  try:
   return {"document":lambda:self.documents.get(s.source_id).original_name,"evidence":lambda:self.evidence.get(s.source_id).title,"timeline":lambda:self.timeline.get(s.source_id).title,"nexus":lambda:self.nexus.get(s.source_id).title,"dbq":lambda:self.dbq.get(s.source_id).title,"rating_strategy":lambda:"Rating Strategy Report","optimizer":lambda:"Claim Optimizer Report","claimant_statement":lambda:self.documents.get(s.source_id).original_name,"witness_statement":lambda:self.documents.get(s.source_id).original_name,"provider_request":lambda:self.documents.get(s.source_id).original_name}[s.source_type]()
  except Exception:return s.display_name or "Unavailable source"
 def _failed(self,pid,eid,status,msg):
  with self.manager._connect() as c:c.execute("UPDATE submission_packages SET status=?,updated_at=? WHERE package_id=?",(status,_now(),pid));c.execute("UPDATE submission_exports SET status=?,error_message=?,completed_at=? WHERE export_id=?",(status,msg,_now(),eid))
 @staticmethod
 def _safe(v):return "".join(c if c.isalnum() or c in "-_." else "-" for c in v).strip("-.") or "package"
