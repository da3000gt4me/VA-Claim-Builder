from __future__ import annotations
import hashlib, json, shutil, zipfile
from pathlib import Path
from typing import Any
from docx import Document

class SubmissionBinderAssembler:
    def assemble(self, project_name: str, claims: list[dict[str, Any]], documents: list[dict[str, Any]],
                 generated_outputs: list[str | Path], output_dir: Path, max_part_mb: int = 95) -> dict[str, Any]:
        output_dir = Path(output_dir); output_dir.mkdir(parents=True, exist_ok=True)
        included = [d for d in documents if d.get('include_in_final')]
        missing = [d.get('original_name', d.get('id')) for d in included if not Path(d.get('original_path','')).exists()]
        outputs = [Path(p) for p in generated_outputs if Path(p).exists()]
        manifest = {
            "project_name": project_name, "claim_count": len(claims), "claims": claims,
            "source_document_count": len(included), "generated_output_count": len(outputs),
            "missing_files": missing, "files": [], "max_part_mb": max_part_mb,
            "status": "blocked" if missing else "assembled_for_human_review",
        }
        if missing: return manifest
        staging = output_dir / "binder_staging"
        if staging.exists(): shutil.rmtree(staging)
        (staging / "01_Generated").mkdir(parents=True); (staging / "02_Evidence").mkdir(parents=True)
        ordered = []
        for i, p in enumerate(outputs, 1): ordered.append((p, staging / "01_Generated" / f"{i:03d}_{p.name}", "generated"))
        for i, d in enumerate(included, 1):
            p=Path(d['original_path']); ordered.append((p, staging / "02_Evidence" / f"{i:03d}_{p.name}", d.get('category','evidence')))
        for src, dest, role in ordered:
            shutil.copy2(src, dest); digest=self._sha(dest)
            manifest['files'].append({"name":str(dest.relative_to(staging)),"role":role,"bytes":dest.stat().st_size,"sha256":digest})
        self._make_index(project_name, claims, manifest, staging / "00_Submission_Index.docx")
        manifest_path=staging/'manifest.json'; manifest_path.write_text(json.dumps(manifest,indent=2,default=str),encoding='utf-8')
        zip_path=output_dir/'VA_Submission_Review_Binder.zip'
        with zipfile.ZipFile(zip_path,'w',zipfile.ZIP_DEFLATED) as z:
            for p in sorted(staging.rglob('*')):
                if p.is_file(): z.write(p,p.relative_to(staging))
        manifest['binder_zip']=str(zip_path); manifest['binder_sha256']=self._sha(zip_path)
        manifest['estimated_parts']=max(1, (zip_path.stat().st_size + max_part_mb*1024*1024 - 1)//(max_part_mb*1024*1024))
        return manifest

    @staticmethod
    def _sha(path: Path) -> str:
        h=hashlib.sha256()
        with path.open('rb') as f:
            for chunk in iter(lambda:f.read(1024*1024),b''): h.update(chunk)
        return h.hexdigest()

    @staticmethod
    def _make_index(project_name, claims, manifest, path):
        doc=Document(); doc.add_heading(f"{project_name} — Submission Review Index",0)
        doc.add_paragraph("Review every item before submission. Draft or unsigned material must not be filed as completed evidence.")
        doc.add_heading("Claimed Conditions",1)
        for c in claims: doc.add_paragraph(f"{c.get('condition_name')} — {c.get('theory','unknown')}",style='List Bullet')
        doc.add_heading("Binder Contents",1)
        for f in manifest['files']: doc.add_paragraph(f"{f['name']} ({f['role']}, {f['bytes']} bytes)",style='List Bullet')
        doc.save(path)
