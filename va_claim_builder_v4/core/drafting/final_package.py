from __future__ import annotations
import json
from datetime import datetime, timezone
from pathlib import Path
from docx import Document

class FinalPackageGenerator:
    """Builds a claim-readiness package from approved evidence and resolved timelines."""
    def __init__(self, output_root: str | Path):
        self.output_root=Path(output_root); self.output_root.mkdir(parents=True, exist_ok=True)

    def generate(self, project_name: str, readiness: list[dict], approved_annotations: list[dict],
                 timeline: list[dict], contradictions: list[dict]) -> tuple[Path, Path]:
        stamp=datetime.now(timezone.utc).strftime('%Y%m%d_%H%M%S')
        docx_path=self.output_root/f"VA_Claim_Readiness_Package_{stamp}.docx"
        json_path=self.output_root/f"VA_Claim_Readiness_Manifest_{stamp}.json"
        doc=Document(); doc.add_heading(project_name,0)
        doc.add_paragraph('PRE-SUBMISSION REVIEW PRODUCT — verify all facts, signatures, forms, and source pages before filing.')
        doc.add_heading('Claim Readiness Dashboard',1)
        table=doc.add_table(rows=1, cols=6)
        for i,h in enumerate(['Condition','Theory','Score','Status','Missing Elements','Open Conflicts']): table.rows[0].cells[i].text=h
        for item in readiness:
            cells=table.add_row().cells
            vals=[item['condition_name'],str(item.get('theory','')),str(item['score']),item['status'],', '.join(item['missing_elements']) or 'None',str(item['contradiction_count'])]
            for i,v in enumerate(vals): cells[i].text=v
        for item in readiness:
            doc.add_page_break(); doc.add_heading(item['condition_name'],0)
            doc.add_paragraph(f"Theory: {item.get('theory')} | Readiness score: {item['score']} | Status: {item['status']}")
            doc.add_heading('Supported Elements',1)
            for x in item['supported_elements'] or ['None approved']: doc.add_paragraph(x,style='List Bullet')
            doc.add_heading('Evidence Gaps',1)
            for x in item['missing_elements'] or ['No required element gaps detected by current rules.']: doc.add_paragraph(x,style='List Bullet')
            doc.add_heading('Approved Evidence',1)
            evidence=[a for a in approved_annotations if a.get('claim_id')==item['claim_id']]
            for a in evidence:
                doc.add_paragraph(f"{a.get('claim_element')}: {a.get('finding')} — {a.get('document_name',a.get('document_id'))}, p. {a.get('page','?')}",style='List Bullet')
            doc.add_heading('Historical Timeline',1)
            events=[e for e in timeline if e.get('claim_id')==item['claim_id']]
            for e in events:
                doc.add_paragraph(f"{e.get('event_date_start') or 'Date unknown'}: {e.get('description')}",style='List Bullet')
            doc.add_heading('Unresolved Contradictions',1)
            conflicts=[c for c in contradictions if c.get('claim_id')==item['claim_id'] and c.get('status','open')=='open']
            for c in conflicts:
                doc.add_paragraph(f"{c.get('summary')} Resolution needed: {c.get('resolution_prompt')}",style='List Bullet')
            if not conflicts: doc.add_paragraph('None open.')
        doc.save(docx_path)
        manifest={'generated_at':datetime.now(timezone.utc).isoformat(),'project_name':project_name,'readiness':readiness,
                  'approved_annotation_ids':[a.get('id') for a in approved_annotations],
                  'timeline_event_ids':[e.get('id') for e in timeline],
                  'open_contradiction_ids':[c.get('id') for c in contradictions if c.get('status','open')=='open']}
        json_path.write_text(json.dumps(manifest,indent=2),encoding='utf-8')
        return docx_path,json_path
