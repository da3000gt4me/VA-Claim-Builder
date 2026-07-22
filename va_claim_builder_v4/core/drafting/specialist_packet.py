from __future__ import annotations
import json
from pathlib import Path
from typing import Any
from docx import Document

SPECIALTY_RULES = {
    "migraine": "Neurology", "headache": "Neurology", "sleep apnea": "Sleep Medicine or Pulmonology",
    "insomnia": "Sleep Medicine or Psychiatry", "anxiety": "Psychiatry or Psychology", "depression": "Psychiatry or Psychology",
    "ibs": "Gastroenterology", "bowel": "Gastroenterology", "rhinitis": "Allergy/Immunology or ENT",
    "sinus": "ENT", "hearing": "Audiology or Otolaryngology", "tinnitus": "Audiology or Otolaryngology",
    "cervical": "Orthopedics, PM&R, or Neurology", "lumbar": "Orthopedics, PM&R, or Neurology",
    "knee": "Orthopedics", "shoulder": "Orthopedics", "tmj": "Dentistry, Oral Surgery, or TMJ Specialist",
    "urinary": "Urology", "voiding": "Urology", "erectile": "Urology", "skin": "Dermatology",
}

class SpecialistPacketGenerator:
    def recommend_specialty(self, condition: str) -> str:
        lower = condition.lower()
        for key, specialty in SPECIALTY_RULES.items():
            if key in lower: return specialty
        return "Appropriately qualified treating clinician or relevant specialist"

    def build(self, claim: dict[str, Any], fact_matrix: dict[str, Any], literature: list[dict[str, Any]] | None = None,
              denial: dict[str, Any] | None = None, dbq: dict[str, Any] | None = None) -> dict[str, Any]:
        condition = claim["condition_name"]
        facts = [f for values in fact_matrix.get("facts_by_element", {}).values() for f in values]
        return {
            "condition_name": condition,
            "theory": claim.get("theory", "unknown"),
            "recommended_specialty": self.recommend_specialty(condition),
            "verified_facts": facts,
            "literature": literature or [],
            "denial_issues": (denial or {}).get("denial_reasons", []),
            "dbq_review_flags": (dbq or {}).get("review_flags", []),
            "clinician_questions": self._questions(claim, fact_matrix, denial, dbq),
            "guardrails": [
                "The clinician must independently verify the record and choose the opinion supported by medical judgment.",
                "General literature does not establish patient-specific causation.",
                "Veteran and witness reports are identified as lay history, not converted into medical findings.",
                "Do not sign until every factual statement and cited record has been checked.",
            ],
        }

    def export(self, packet: dict[str, Any], output_dir: Path) -> list[Path]:
        output_dir = Path(output_dir); output_dir.mkdir(parents=True, exist_ok=True)
        safe = ''.join(c if c.isalnum() else '_' for c in packet['condition_name']).strip('_')
        docx_path = output_dir / f"{safe}_Specialist_Nexus_Review_Packet.docx"
        json_path = output_dir / f"{safe}_Specialist_Nexus_Review_Packet.json"
        doc = Document(); doc.add_heading(f"Specialist Nexus Review Packet — {packet['condition_name']}", 0)
        doc.add_paragraph(f"Suggested specialty: {packet['recommended_specialty']}")
        doc.add_paragraph(f"Claim theory: {packet['theory']}")
        doc.add_heading("Independent Review Requirements", 1)
        for g in packet['guardrails']: doc.add_paragraph(g, style='List Bullet')
        doc.add_heading("Verified Evidence for Review", 1)
        if not packet['verified_facts']: doc.add_paragraph("No approved draftable facts are available.")
        for f in packet['verified_facts']:
            doc.add_paragraph(f.get('proposition') or f.get('finding') or '', style='List Bullet')
            src = f.get('source') or f.get('source_label') or f.get('document_name')
            if src: doc.add_paragraph(f"Source: {src}")
        doc.add_heading("Prior Denial / Examination Issues", 1)
        for x in packet['denial_issues'] + packet['dbq_review_flags'] or ["None supplied."]:
            doc.add_paragraph(str(x), style='List Bullet')
        doc.add_heading("Medical Literature for Independent Consideration", 1)
        for item in packet['literature'] or [{"title":"None supplied."}]:
            doc.add_paragraph(item.get('title','Untitled'), style='List Bullet')
            doc.add_paragraph(item.get('allowed_use',''))
        doc.add_heading("Questions for the Clinician", 1)
        for q in packet['clinician_questions']: doc.add_paragraph(q, style='List Number')
        doc.add_heading("Clinician Opinion and Rationale", 1)
        doc.add_paragraph("Opinion selected: ______________________________________________")
        doc.add_paragraph("Medical rationale, including positive and negative evidence and competing causes:\n\n\n")
        doc.add_paragraph("Records reviewed: ______________________________________________")
        doc.add_paragraph("Name / credentials / signature / date: __________________________")
        doc.save(docx_path)
        json_path.write_text(json.dumps(packet, indent=2, default=str), encoding='utf-8')
        return [docx_path, json_path]

    @staticmethod
    def _questions(claim, matrix, denial, dbq):
        condition = claim['condition_name']; theory = claim.get('theory','unknown')
        qs = [f"What is the current diagnosis and medically supported severity of {condition}?",
              f"After reviewing the cited evidence, is {condition} at least as likely as not related under the {theory} theory?",
              "What medical reasoning connects—or does not connect—the verified facts to the opinion?",
              "What competing causes were considered, and why are they more or less persuasive?"]
        if denial and denial.get('denial_reasons'): qs.append("How does your opinion address the prior VA denial reason(s)?")
        if dbq and dbq.get('review_flags'): qs.append("Can the missing or unclear DBQ/C&P findings be clarified?")
        return qs
