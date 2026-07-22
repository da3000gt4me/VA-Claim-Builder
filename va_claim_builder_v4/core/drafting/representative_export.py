from __future__ import annotations
from pathlib import Path
import json
from docx import Document

class AccreditedRepresentativeExport:
    def build(self, claims: list[dict], theory_results: list[dict], adversarial_results: list[dict], denial_issues: list[dict], legal_map: list[dict]) -> dict:
        return {"claims": claims, "theory_comparisons": theory_results, "adversarial_reviews": adversarial_results, "denial_issues": denial_issues, "legal_issue_map": legal_map, "notice": "Prepared for review by a VA-accredited representative. This export is not legal advice or representation."}
    def export(self, package: dict, output_dir: Path) -> tuple[Path, Path]:
        output_dir.mkdir(parents=True, exist_ok=True)
        json_path=output_dir/"Accredited_Representative_Review_Export.json"; json_path.write_text(json.dumps(package, indent=2, default=str),encoding="utf-8")
        docx_path=output_dir/"Accredited_Representative_Review_Export.docx"
        doc=Document(); doc.add_heading("VA Claim Builder — Accredited Representative Review Export",0); doc.add_paragraph(package["notice"])
        for claim in package["claims"]:
            doc.add_heading(claim.get("condition_name","Claim"), level=1)
            theory=next((x for x in package["theory_comparisons"] if x.get("claim_id")==claim.get("id")),{})
            doc.add_paragraph(f"Current theory: {claim.get('theory','unknown')}; evidence-coverage comparison suggests: {theory.get('recommended_primary_theory') or 'no supported theory established'}.")
            review=next((x for x in package["adversarial_reviews"] if x.get("claim_id")==claim.get("id")),{})
            doc.add_paragraph(f"Defensibility score: {review.get('defensibility_score','N/A')}")
            for finding in review.get("findings",[]): doc.add_paragraph(f"[{finding.get('severity','').upper()}] {finding.get('issue')} — {finding.get('recommended_fix')}",style="List Bullet")
        doc.save(docx_path); return docx_path,json_path
