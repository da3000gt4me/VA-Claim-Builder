from __future__ import annotations
import json,sqlite3,uuid
from dataclasses import dataclass
from datetime import datetime,timezone
from pathlib import Path
from typing import Any
from core.projects.manager import DBQ_SCHEMA
from .templates import get_template
def _now():return datetime.now(timezone.utc).isoformat()
DBQ_STATUSES=("draft","pending","completed","failed","cancelled","reviewed")
@dataclass(frozen=True,slots=True)
class DBQRevision:revision_id:str;dbq_id:str;revision_number:int;fields:dict[str,str];suggestions:list[dict[str,Any]];ai_metadata:dict[str,Any];created_at:str
@dataclass(frozen=True,slots=True)
class DBQRecord:
 dbq_id:str;claim_id:str;template_id:str;status:str;title:str;condition:str;exam_date:str;examiner_name:str;examiner_credentials:str;specialty:str;revision_number:int;created_at:str;updated_at:str;fields:dict[str,str];suggestions:tuple[dict[str,Any],...];evidence_ids:tuple[str,...];timeline_event_ids:tuple[str,...];document_ids:tuple[str,...]
class DBQManager:
 def __init__(self,project):self.project=project;self._ensure()
 def _connect(self):c=sqlite3.connect(self.project.database_path);c.row_factory=sqlite3.Row;c.execute("PRAGMA foreign_keys=ON");return c
 def create(self,claim_id,template_id,title,*,status="draft",condition="",exam_date="",examiner_name="",examiner_credentials="",specialty="",fields=None,evidence_ids=(),timeline_event_ids=(),document_ids=()):
  get_template(template_id);self._validate(status,title);did=str(uuid.uuid4());now=_now();values=self._fields(template_id,fields)
  with self._connect() as c:c.execute("INSERT INTO dbq_records VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?)",(did,claim_id,template_id,status,title.strip(),condition,exam_date,examiner_name,examiner_credentials,specialty,1,now,now));self._revision(c,did,1,values,[],{})
  for x in evidence_ids:self.link_evidence(did,x)
  for x in timeline_event_ids:self.link_timeline_event(did,x)
  for x in document_ids:self.link_document(did,x)
  return self.get(did)
 def get(self,did):
  with self._connect() as c:
   r=c.execute("SELECT * FROM dbq_records WHERE dbq_id=?",(did,)).fetchone()
   if not r:raise KeyError(did)
   v=c.execute("SELECT * FROM dbq_revisions WHERE dbq_id=? ORDER BY revision_number DESC LIMIT 1",(did,)).fetchone();rel=lambda t,col:tuple(x[0] for x in c.execute(f"SELECT {col} FROM {t} WHERE dbq_id=? ORDER BY {col}",(did,)))
   return DBQRecord(*tuple(r),json.loads(v[3]),tuple(json.loads(v[4])),rel("dbq_evidence","evidence_id"),rel("dbq_timeline","event_id"),rel("dbq_documents","document_id"))
 def update(self,did,**changes):
  old=self.get(did);meta={k:changes.pop(k) for k in list(changes) if k in {"template_id","status","title","condition","exam_date","examiner_name","examiner_credentials","specialty"}};template=meta.get("template_id",old.template_id);get_template(template);self._validate(meta.get("status",old.status),meta.get("title",old.title));fields=self._fields(template,{**old.fields,**changes});num=old.revision_number+1;vals={k:getattr(old,k) for k in ("template_id","status","title","condition","exam_date","examiner_name","examiner_credentials","specialty")};vals.update(meta)
  with self._connect() as c:c.execute("UPDATE dbq_records SET template_id=?,status=?,title=?,condition=?,exam_date=?,examiner_name=?,examiner_credentials=?,specialty=?,revision_number=?,updated_at=? WHERE dbq_id=?",(*(vals[k] for k in ("template_id","status","title","condition","exam_date","examiner_name","examiner_credentials","specialty")),num,_now(),did));self._revision(c,did,num,fields,list(old.suggestions),{})
  return self.get(did)
 def create_ai_revision(self,did,suggestions,metadata):
  old=self.get(did);num=old.revision_number+1
  with self._connect() as c:c.execute("UPDATE dbq_records SET revision_number=?,status='completed',updated_at=? WHERE dbq_id=?",(num,_now(),did));self._revision(c,did,num,old.fields,suggestions,metadata)
  return self.get(did)
 def decide_suggestion(self,did,field_key,accept):
  old=self.get(did);suggestions=[];fields=dict(old.fields)
  for s in old.suggestions:
   s=dict(s)
   if s["field_key"]==field_key:s["decision"]="accepted" if accept else "rejected";fields[field_key]=s["proposed_value"] if accept else fields.get(field_key,"")
   suggestions.append(s)
  num=old.revision_number+1
  with self._connect() as c:c.execute("UPDATE dbq_records SET revision_number=?,updated_at=? WHERE dbq_id=?",(num,_now(),did));self._revision(c,did,num,fields,suggestions,{})
  return self.get(did)
 def revisions(self,did):
  with self._connect() as c:rows=c.execute("SELECT * FROM dbq_revisions WHERE dbq_id=? ORDER BY revision_number DESC",(did,)).fetchall()
  return [DBQRevision(r[0],r[1],r[2],json.loads(r[3]),json.loads(r[4]),json.loads(r[5]),r[6]) for r in rows]
 def list(self,*,claim_id=None,template_id=None,status=None,search=""):
  q="SELECT dbq_id FROM dbq_records WHERE 1=1";a=[]
  for col,val in (("claim_id",claim_id),("template_id",template_id),("status",status)):
   if val:q+=f" AND {col}=?";a.append(val)
  if search:q+=" AND (title LIKE ? OR condition LIKE ? OR examiner_name LIKE ?)";a += [f"%{search}%"]*3
  q+=" ORDER BY updated_at DESC,dbq_id";
  with self._connect() as c:ids=[r[0] for r in c.execute(q,a)]
  return [self.get(x) for x in ids]
 def duplicate(self,did):
  x=self.get(did);return self.create(x.claim_id,x.template_id,"Copy of "+x.title,status="draft",condition=x.condition,exam_date=x.exam_date,examiner_name=x.examiner_name,examiner_credentials=x.examiner_credentials,specialty=x.specialty,fields=x.fields,evidence_ids=x.evidence_ids,timeline_event_ids=x.timeline_event_ids,document_ids=x.document_ids)
 def delete(self,did):
  with self._connect() as c:
   if c.execute("DELETE FROM dbq_records WHERE dbq_id=?",(did,)).rowcount==0:raise KeyError(did)
 def _link(self,did,x,table,col):
  with self._connect() as c:c.execute(f"INSERT OR IGNORE INTO {table}(dbq_id,{col},linked_at) VALUES(?,?,?)",(did,x,_now()))
 def _unlink(self,did,x,table,col):
  with self._connect() as c:c.execute(f"DELETE FROM {table} WHERE dbq_id=? AND {col}=?",(did,x))
 def link_evidence(self,d,x):self._link(d,x,"dbq_evidence","evidence_id")
 def unlink_evidence(self,d,x):self._unlink(d,x,"dbq_evidence","evidence_id")
 def link_timeline_event(self,d,x):self._link(d,x,"dbq_timeline","event_id")
 def unlink_timeline_event(self,d,x):self._unlink(d,x,"dbq_timeline","event_id")
 def link_document(self,d,x):self._link(d,x,"dbq_documents","document_id")
 def unlink_document(self,d,x):self._unlink(d,x,"dbq_documents","document_id")
 def completeness(self,did):
  x=self.get(did);t=get_template(x.template_id);required=[f for f in t.fields if f.required and f.kind!="examiner"];completed=[f.key for f in required if x.fields.get(f.key,"").strip()];incomplete=[f.key for f in required if f.key not in completed];examiner=[f.key for f in t.fields if f.kind=="examiner"];conflicts=[s["field_key"] for s in x.suggestions if s.get("conflicting_information")];gaps=list(dict.fromkeys(incomplete+[g for s in x.suggestions for g in s.get("missing_information",[])]));return {"score":round(100*len(completed)/len(required)) if required else 100,"completed":completed,"incomplete":incomplete,"conflicting":conflicts,"examiner_only":examiner,"supporting_evidence_count":len(x.evidence_ids)+len(x.timeline_event_ids)+len(x.document_ids),"evidence_gaps":gaps,"disclaimer":"Completeness does not guarantee a rating or favorable decision."}
 def export_docx(self,did,path):
  from docx import Document
  x=self.get(did);t=get_template(x.template_id);gap=self.completeness(did);d=Document();d.add_heading(f"DBQ Preparation Packet — {t.name}",0);d.add_paragraph("This is not an official VA DBQ. It requires completion and review by an appropriate examiner or provider.");d.add_paragraph(f"Condition / claim: {x.condition or self._claim(x.claim_id)}");sections=[]
  for f in t.fields:
   if f.section not in sections:sections.append(f.section);d.add_heading(f.section,1)
   d.add_heading(f.label,2);d.add_paragraph(x.fields.get(f.key,"") or "Missing");d.add_paragraph(f"Information class: {f.kind.replace('_',' ')}"+(" — Examiner/provider completion required." if f.kind=="examiner" else ""))
   for s in x.suggestions:
    if s["field_key"]==f.key:d.add_paragraph("Traceability: "+", ".join(s.get("source_labels",[])));d.add_paragraph("Unresolved conflict: "+s.get("conflicting_information","") if s.get("conflicting_information") else "")
  d.add_heading("Completeness and Gaps",1);d.add_paragraph(f"Completeness: {gap['score']}%. Missing: {', '.join(gap['incomplete']) or 'None'}. Examiner-only: {', '.join(gap['examiner_only']) or 'None'}.");p=Path(path);p.parent.mkdir(parents=True,exist_ok=True);d.save(p);return p
 def _claim(self,cid):
  with self._connect() as c:r=c.execute("SELECT condition_name FROM claims WHERE claim_id=?",(cid,)).fetchone();return r[0] if r else ""
 def _ensure(self):
  with self._connect() as c:c.executescript(DBQ_SCHEMA)
 def _revision(self,c,did,num,fields,suggestions,meta):c.execute("INSERT INTO dbq_revisions VALUES(?,?,?,?,?,?,?)",(str(uuid.uuid4()),did,num,json.dumps(fields),json.dumps(suggestions),json.dumps(meta),_now()))
 @staticmethod
 def _fields(template_id,values):return {f.key:str((values or {}).get(f.key,"")) for f in get_template(template_id).fields}
 @staticmethod
 def _validate(status,title):
  if status not in DBQ_STATUSES:raise ValueError("Unsupported DBQ status")
  if not str(title).strip():raise ValueError("DBQ title is required")
