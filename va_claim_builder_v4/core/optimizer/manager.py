from __future__ import annotations
import json,sqlite3,uuid
from dataclasses import dataclass
from datetime import datetime,timezone
from pathlib import Path
from typing import Any
from core.projects.manager import OPTIMIZER_SCHEMA
from .taxonomy import GAP_CATEGORIES
def _now():return datetime.now(timezone.utc).isoformat()
@dataclass(frozen=True,slots=True)
class OptimizerAssessment:assessment_id:str;claim_id:str;job_id:str|None;status:str;assessment_timestamp:str;assessment_version:str;overall_score:int;service_connection_score:int;severity_rating_score:int;evidence_quality_score:int;evidence_consistency_score:int;confidence:str;score_explanation:dict[str,Any];ai_metadata:dict[str,Any];error_message:str;created_at:str;updated_at:str
@dataclass(frozen=True,slots=True)
class OptimizerGap:gap_id:str;assessment_id:str;category:str;description:str;severity:str;priority:int;confidence:str;evidence_basis:tuple[str,...];recommended_resolution:str;responsible_party:str;status:str;follow_up_date:str;notes:str;origin:str;user_confirmation:str;created_at:str;updated_at:str
@dataclass(frozen=True,slots=True)
class OptimizerAction:action_id:str;assessment_id:str;gap_id:str|None;description:str;priority:int;expected_value:str;effort_level:str;dependency:str;status:str;completion_notes:str;origin:str;created_at:str;updated_at:str
class OptimizerManager:
 def __init__(self,project):self.project=project;self._ensure()
 def _connect(self):c=sqlite3.connect(self.project.database_path);c.row_factory=sqlite3.Row;c.execute("PRAGMA foreign_keys=ON");return c
 def create(self,claim_id,*,job_id=None,status="pending",assessment_version="1.0",overall_score=0,service_connection_score=0,severity_rating_score=0,evidence_quality_score=0,evidence_consistency_score=0,confidence="low",score_explanation=None,ai_metadata=None,error_message=""):
  aid=str(uuid.uuid4());now=_now()
  with self._connect() as c:c.execute("INSERT INTO optimizer_assessments VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)",(aid,claim_id,job_id,status,now,assessment_version,self._score(overall_score),self._score(service_connection_score),self._score(severity_rating_score),self._score(evidence_quality_score),self._score(evidence_consistency_score),confidence,json.dumps(score_explanation or {}),json.dumps(ai_metadata or {}),error_message,now,now))
  return self.get(aid)
 def get(self,aid):
  with self._connect() as c:r=c.execute("SELECT * FROM optimizer_assessments WHERE assessment_id=?",(aid,)).fetchone()
  if not r:raise KeyError(aid)
  return OptimizerAssessment(r[0],r[1],r[2],r[3],r[4],r[5],*map(int,r[6:11]),r[11],json.loads(r[12]),json.loads(r[13]),r[14],r[15],r[16])
 def update(self,aid,**changes):
  allowed={"status","overall_score","service_connection_score","severity_rating_score","evidence_quality_score","evidence_consistency_score","confidence","score_explanation","ai_metadata","error_message"};unknown=set(changes)-allowed
  if unknown:raise ValueError("Unsupported assessment fields")
  cols=[];args=[]
  for k,v in changes.items():cols.append((k+"_json" if k in {"score_explanation","ai_metadata"} else k)+"=?");args.append(json.dumps(v) if k in {"score_explanation","ai_metadata"} else self._score(v) if k.endswith("score") else v)
  with self._connect() as c:c.execute("UPDATE optimizer_assessments SET "+",".join(cols)+",updated_at=? WHERE assessment_id=?",(*args,_now(),aid))
  return self.get(aid)
 def delete(self,aid):
  with self._connect() as c:
   if c.execute("DELETE FROM optimizer_assessments WHERE assessment_id=?",(aid,)).rowcount==0:raise KeyError(aid)
 def list(self,*,claim_id=None,status=None,confidence=None,search=""):
  q="SELECT DISTINCT a.* FROM optimizer_assessments a LEFT JOIN optimizer_gaps g ON g.assessment_id=a.assessment_id WHERE 1=1";args=[]
  for col,val in (("a.claim_id",claim_id),("a.status",status),("a.confidence",confidence)):
   if val:q+=f" AND {col}=?";args.append(val)
  if search:q+=" AND (g.description LIKE ? OR g.recommended_resolution LIKE ? OR a.score_explanation_json LIKE ?)";args += [f"%{search}%"]*3
  q+=" ORDER BY a.assessment_timestamp DESC,a.assessment_id"
  with self._connect() as c:rows=c.execute(q,args).fetchall()
  return [self.get(r[0]) for r in rows]
 def history(self,cid):return self.list(claim_id=cid)
 def add_gap(self,aid,category,description,*,severity="medium",priority=3,confidence="medium",evidence_basis=(),recommended_resolution="",responsible_party="",status="unresolved",follow_up_date="",notes="",origin="manual",user_confirmation="confirmed",preserve=True):
  if category not in GAP_CATEGORIES:raise ValueError("Unsupported gap category")
  a=self.get(aid);prior=self._prior_decision(a.claim_id,category,aid) if preserve else None
  if prior:status=prior.status;responsible_party=prior.responsible_party;follow_up_date=prior.follow_up_date;notes=prior.notes;user_confirmation=prior.user_confirmation
  gid=str(uuid.uuid4());now=_now()
  with self._connect() as c:c.execute("INSERT INTO optimizer_gaps VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)",(gid,aid,category,description,severity,int(priority),confidence,json.dumps(list(evidence_basis)),recommended_resolution,responsible_party,status,follow_up_date,notes,origin,user_confirmation,now,now))
  return self.get_gap(gid)
 def get_gap(self,gid):
  with self._connect() as c:r=c.execute("SELECT * FROM optimizer_gaps WHERE gap_id=?",(gid,)).fetchone()
  if not r:raise KeyError(gid)
  return OptimizerGap(r[0],r[1],r[2],r[3],r[4],r[5],r[6],tuple(json.loads(r[7])),*r[8:])
 def update_gap(self,gid,**changes):
  allowed={"category","description","severity","priority","confidence","evidence_basis","recommended_resolution","responsible_party","status","follow_up_date","notes","user_confirmation"};unknown=set(changes)-allowed
  if unknown:raise ValueError("Unsupported gap fields")
  if "category" in changes and changes["category"] not in GAP_CATEGORIES:raise ValueError("Unsupported gap category")
  cols=[];args=[]
  for k,v in changes.items():cols.append((k+"_json" if k=="evidence_basis" else k)+"=?");args.append(json.dumps(list(v)) if k=="evidence_basis" else v)
  with self._connect() as c:c.execute("UPDATE optimizer_gaps SET "+",".join(cols)+",updated_at=? WHERE gap_id=?",(*args,_now(),gid))
  return self.get_gap(gid)
 def resolve_gap(self,gid,notes=""):return self.update_gap(gid,status="resolved",notes=notes or self.get_gap(gid).notes,user_confirmation="confirmed")
 def reopen_gap(self,gid):return self.update_gap(gid,status="unresolved")
 def not_applicable(self,gid,notes=""):return self.update_gap(gid,status="not_applicable",notes=notes,user_confirmation="confirmed")
 def reject_gap(self,gid,notes=""):return self.update_gap(gid,status="rejected",notes=notes,user_confirmation="rejected")
 def gaps(self,aid,*,priority=None,category=None,status=None,responsible_party=None,unresolved_only=False,search=""):
  q="SELECT * FROM optimizer_gaps WHERE assessment_id=?";a=[aid]
  for col,val in (("priority",priority),("category",category),("status",status),("responsible_party",responsible_party)):
   if val not in (None,""):q+=f" AND {col}=?";a.append(val)
  if unresolved_only:q+=" AND status='unresolved'"
  if search:q+=" AND (description LIKE ? OR recommended_resolution LIKE ? OR notes LIKE ? OR evidence_basis_json LIKE ?)";a += [f"%{search}%"]*4
  q+=" ORDER BY priority,severity DESC,created_at"
  with self._connect() as c:rows=c.execute(q,a).fetchall()
  return [self.get_gap(r[0]) for r in rows]
 def add_action(self,aid,description,*,gap_id=None,priority=3,expected_value="medium",effort_level="medium",dependency="",status="pending",completion_notes="",origin="rule"):
  action=str(uuid.uuid4());now=_now()
  with self._connect() as c:c.execute("INSERT INTO optimizer_actions VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?)",(action,aid,gap_id,description,int(priority),expected_value,effort_level,dependency,status,completion_notes,origin,now,now))
  return self.get_action(action)
 def get_action(self,x):
  with self._connect() as c:r=c.execute("SELECT * FROM optimizer_actions WHERE action_id=?",(x,)).fetchone()
  if not r:raise KeyError(x)
  return OptimizerAction(*r)
 def actions(self,aid,status=None):
  q="SELECT * FROM optimizer_actions WHERE assessment_id=?";a=[aid]
  if status:q+=" AND status=?";a.append(status)
  q+=" ORDER BY priority,CASE expected_value WHEN 'high' THEN 0 WHEN 'medium' THEN 1 ELSE 2 END,created_at"
  with self._connect() as c:return [OptimizerAction(*r) for r in c.execute(q,a)]
 def update_action(self,action_id,**changes):
  allowed={"priority","expected_value","effort_level","dependency","status","completion_notes"}
  if set(changes)-allowed:raise ValueError("Unsupported action fields")
  with self._connect() as c:c.execute("UPDATE optimizer_actions SET "+",".join(f"{k}=?" for k in changes)+",updated_at=? WHERE action_id=?",(*changes.values(),_now(),action_id))
  return self.get_action(action_id)
 def export_lay_statement(self,aid,path,witness_type="claimant"):
  from docx import Document
  a=self.get(aid);claim=self._claim(a.claim_id);d=Document();d.add_heading(f"Draft {witness_type.title()} Statement Outline",0);d.add_paragraph("DRAFT — Review and signature are required from the actual witness. Include only facts personally known to the witness.");d.add_paragraph(f"Claimed condition: {claim['condition_name']}");d.add_heading("Facts the witness may personally know",1)
  facts=[claim["symptoms"],claim["service_event"],claim["description"]]
  for f in facts:
   if f:d.add_paragraph(f,style="List Bullet")
  d.add_heading("Observation prompts",1);d.add_paragraph("When did you personally observe the symptoms? How did they affect daily activities or work? What changes did you personally observe? Do not guess about diagnosis or causation.");d.add_paragraph("Witness name: ____________________   Signature: ____________________   Date: __________");p=Path(path);p.parent.mkdir(parents=True,exist_ok=True);d.save(p);return p
 def export_provider_request(self,aid,path):
  from docx import Document
  a=self.get(aid);claim=self._claim(a.claim_id);gaps=[g for g in self.gaps(aid,unresolved_only=True) if g.category in {"missing_current_diagnosis","weak_nexus_evidence","missing_secondary_causation","missing_secondary_aggravation","missing_preexisting_aggravation","missing_examiner_findings","missing_testing_imaging","missing_provider_statement"}];d=Document();d.add_heading("Provider Review Request Packet",0);d.add_paragraph("This request does not ask or require a favorable conclusion. Please provide an independent medical assessment.");d.add_paragraph(f"Claimed condition: {claim['condition_name']}");d.add_heading("Specific medical questions and requested clarification",1)
  for g in gaps:d.add_paragraph(g.recommended_resolution or g.description,style="List Bullet");d.add_paragraph("Relevant records summary: "+("; ".join(g.evidence_basis) or "See selected project records."))
  d.add_heading("Provider completion",1);d.add_paragraph("Findings and rationale: ____________________________________________________");d.add_paragraph("Provider name/signature: ____________________  Credentials: ______________  Date: ________");p=Path(path);p.parent.mkdir(parents=True,exist_ok=True);d.save(p);return p
 def _prior_decision(self,cid,category,current):
  with self._connect() as c:r=c.execute("SELECT g.gap_id FROM optimizer_gaps g JOIN optimizer_assessments a ON a.assessment_id=g.assessment_id WHERE a.claim_id=? AND g.category=? AND g.assessment_id<>? AND (g.user_confirmation<>'pending' OR g.status<>'unresolved') ORDER BY a.assessment_timestamp DESC LIMIT 1",(cid,category,current)).fetchone()
  return self.get_gap(r[0]) if r else None
 def _claim(self,cid):
  with self._connect() as c:r=c.execute("SELECT condition_name,description,service_event,symptoms FROM claims WHERE claim_id=?",(cid,)).fetchone();return dict(r) if r else {"condition_name":"","description":"","service_event":"","symptoms":""}
 def _ensure(self):
  with self._connect() as c:c.executescript(OPTIMIZER_SCHEMA)
 @staticmethod
 def _score(v):return max(0,min(100,int(v)))
