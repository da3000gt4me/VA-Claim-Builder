from __future__ import annotations
import json,sqlite3,uuid
from dataclasses import dataclass
from datetime import datetime,timezone
from pathlib import Path
from typing import Any
from core.projects import ProjectInfo
from core.projects.manager import NEXUS_SCHEMA

def _now(): return datetime.now(timezone.utc).isoformat()
NEXUS_STATUSES=("draft","pending","completed","failed","cancelled","final")
NEXUS_THEORIES=("direct","secondary","preexisting_aggravation","secondary_aggravation","presumptive","continuity")
CONTENT_FIELDS=("records_reviewed","service_history","current_diagnosis","medical_history","in_service_event","nexus_opinion","medical_rationale","favorable_evidence","unfavorable_evidence","probability_language","limitations","signature_block","user_facts","missing_facts")

@dataclass(frozen=True,slots=True)
class NexusVersion:
 version_id:str;letter_id:str;version_number:int;content:dict[str,Any];ai_metadata:dict[str,Any];created_at:str
@dataclass(frozen=True,slots=True)
class NexusLetter:
 letter_id:str;claim_id:str;title:str;letter_type:str;status:str;primary_theory:str;theories:tuple[str,...];author_name:str;author_credentials:str;specialty:str;current_version:int;created_at:str;updated_at:str;content:dict[str,Any];evidence_ids:tuple[str,...];timeline_event_ids:tuple[str,...];document_ids:tuple[str,...]

class NexusLetterManager:
 def __init__(self,project:ProjectInfo):self.project=project;self._ensure()
 def _connect(self):c=sqlite3.connect(self.project.database_path);c.row_factory=sqlite3.Row;c.execute("PRAGMA foreign_keys=ON");return c
 def create(self,claim_id,title,*,letter_type="physician_review",status="draft",primary_theory="direct",theories=(),author_name="",author_credentials="",specialty="",content=None,evidence_ids=(),timeline_event_ids=(),document_ids=()):
  if not str(title).strip():raise ValueError("Nexus letter title is required")
  self._validate(status,primary_theory,theories);lid=str(uuid.uuid4());now=_now();all_theories=list(dict.fromkeys([primary_theory,*theories]))
  with self._connect() as c:
   c.execute("INSERT INTO nexus_letters VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?)",(lid,claim_id,title.strip(),letter_type,status,primary_theory,json.dumps(all_theories),author_name,author_credentials,specialty,1,now,now))
   self._version(c,lid,1,self._content(content),{})
  for x in evidence_ids:self.link_evidence(lid,x)
  for x in timeline_event_ids:self.link_timeline_event(lid,x)
  for x in document_ids:self.link_document(lid,x)
  return self.get(lid)
 def get(self,lid):
  with self._connect() as c:
   r=c.execute("SELECT * FROM nexus_letters WHERE letter_id=?",(lid,)).fetchone()
   if not r:raise KeyError(lid)
   v=c.execute("SELECT * FROM nexus_letter_versions WHERE letter_id=? ORDER BY version_number DESC LIMIT 1",(lid,)).fetchone()
   rel=lambda t,col:tuple(x[0] for x in c.execute(f"SELECT {col} FROM {t} WHERE letter_id=? ORDER BY {col}",(lid,)))
   return NexusLetter(r[0],r[1],r[2],r[3],r[4],r[5],tuple(json.loads(r[6])),r[7],r[8],r[9],r[10],r[11],r[12],json.loads(v[3]),rel("nexus_letter_evidence","evidence_id"),rel("nexus_letter_timeline","event_id"),rel("nexus_letter_documents","document_id"))
 def update(self,lid,**changes):
  old=self.get(lid);meta={k:changes.pop(k) for k in list(changes) if k in {"title","letter_type","status","primary_theory","theories","author_name","author_credentials","specialty"}}
  content={**old.content,**changes};status=meta.get("status",old.status);primary=meta.get("primary_theory",old.primary_theory);theories=meta.get("theories",old.theories);self._validate(status,primary,theories);num=old.current_version+1;now=_now()
  values={"title":old.title,"letter_type":old.letter_type,"status":old.status,"primary_theory":old.primary_theory,"theories":old.theories,"author_name":old.author_name,"author_credentials":old.author_credentials,"specialty":old.specialty};values.update(meta)
  if not str(values["title"]).strip():raise ValueError("Nexus letter title is required")
  with self._connect() as c:
   c.execute("UPDATE nexus_letters SET title=?,letter_type=?,status=?,primary_theory=?,theories_json=?,author_name=?,author_credentials=?,specialty=?,current_version=?,updated_at=? WHERE letter_id=?",(values["title"],values["letter_type"],values["status"],values["primary_theory"],json.dumps(list(dict.fromkeys([values["primary_theory"],*values["theories"]]))),values["author_name"],values["author_credentials"],values["specialty"],num,now,lid));self._version(c,lid,num,self._content(content),{})
  return self.get(lid)
 def create_ai_version(self,lid,content,metadata):
  old=self.get(lid);num=old.current_version+1
  with self._connect() as c:c.execute("UPDATE nexus_letters SET current_version=?,status='completed',updated_at=? WHERE letter_id=?",(num,_now(),lid));self._version(c,lid,num,self._content(content),metadata)
  return self.get(lid)
 def versions(self,lid):
  with self._connect() as c:rows=c.execute("SELECT * FROM nexus_letter_versions WHERE letter_id=? ORDER BY version_number DESC",(lid,)).fetchall()
  return [NexusVersion(r[0],r[1],r[2],json.loads(r[3]),json.loads(r[4]),r[5]) for r in rows]
 def list(self,*,claim_id=None,status=None,theory=None,search=""):
  q="SELECT letter_id FROM nexus_letters WHERE 1=1";a=[]
  for col,val in (("claim_id",claim_id),("status",status),("primary_theory",theory)):
   if val:q+=f" AND {col}=?";a.append(val)
  if search:q+=" AND (title LIKE ? OR author_name LIKE ?)";a += [f"%{search}%"]*2
  q+=" ORDER BY updated_at DESC,letter_id"
  with self._connect() as c:ids=[r[0] for r in c.execute(q,a)]
  return [self.get(x) for x in ids]
 def duplicate(self,lid):
  x=self.get(lid);return self.create(x.claim_id,f"Copy of {x.title}",letter_type=x.letter_type,status="draft",primary_theory=x.primary_theory,theories=x.theories,author_name=x.author_name,author_credentials=x.author_credentials,specialty=x.specialty,content=x.content,evidence_ids=x.evidence_ids,timeline_event_ids=x.timeline_event_ids,document_ids=x.document_ids)
 def delete(self,lid):
  with self._connect() as c:
   if c.execute("DELETE FROM nexus_letters WHERE letter_id=?",(lid,)).rowcount==0:raise KeyError(lid)
 def _link(self,lid,value,table,col):
  with self._connect() as c:c.execute(f"INSERT OR IGNORE INTO {table}(letter_id,{col},linked_at) VALUES(?,?,?)",(lid,value,_now()))
 def _unlink(self,lid,value,table,col):
  with self._connect() as c:c.execute(f"DELETE FROM {table} WHERE letter_id=? AND {col}=?",(lid,value))
 def link_evidence(self,lid,x):self._link(lid,x,"nexus_letter_evidence","evidence_id")
 def unlink_evidence(self,lid,x):self._unlink(lid,x,"nexus_letter_evidence","evidence_id")
 def link_timeline_event(self,lid,x):self._link(lid,x,"nexus_letter_timeline","event_id")
 def unlink_timeline_event(self,lid,x):self._unlink(lid,x,"nexus_letter_timeline","event_id")
 def link_document(self,lid,x):self._link(lid,x,"nexus_letter_documents","document_id")
 def unlink_document(self,lid,x):self._unlink(lid,x,"nexus_letter_documents","document_id")
 def add_literature(self,lid,title,*,author="",publication="",year="",doi_url="",relevance_note="",verified=True):
  rid=str(uuid.uuid4())
  with self._connect() as c:c.execute("INSERT INTO nexus_literature VALUES(?,?,?,?,?,?,?,?,?,?)",(rid,lid,title,author,publication,year,doi_url,relevance_note,int(verified),_now()))
  return rid
 def literature(self,lid):
  with self._connect() as c:return [dict(r) for r in c.execute("SELECT * FROM nexus_literature WHERE letter_id=? ORDER BY year,title",(lid,))]
 def export_docx(self,lid,path,*,claimant_name="",include_claimant=False):
  from docx import Document
  x=self.get(lid);d=Document();d.add_heading(x.title,0);d.add_paragraph("DRAFT — Requires review and signature by a qualified medical professional.")
  if include_claimant and claimant_name.strip():d.add_paragraph(f"Claimant: {claimant_name.strip()}")
  d.add_paragraph(f"Claim/condition: {self._claim_name(x.claim_id)}");d.add_paragraph(f"Primary nexus theory: {x.primary_theory.replace('_',' ').title()}")
  labels=(("records_reviewed","Records Reviewed"),("current_diagnosis","Current Diagnosis"),("service_history","Service History"),("medical_history","Medical History"),("in_service_event","In-Service Event, Injury, Exposure, or Aggravation"),("nexus_opinion","Nexus Opinion"),("medical_rationale","Medical Rationale"),("favorable_evidence","Favorable Evidence"),("unfavorable_evidence","Unfavorable Evidence"),("limitations","Limitations and Caveats"),("signature_block","Provider Signature"))
  for key,label in labels:d.add_heading(label,level=1);d.add_paragraph(str(x.content.get(key,"") or "Not provided."))
  d.add_heading("Literature References",level=1)
  for r in self.literature(lid):d.add_paragraph(f"{r['author']} ({r['year']}). {r['title']}. {r['publication']}. {r['doi_url']}".strip())
  p=Path(path);p.parent.mkdir(parents=True,exist_ok=True);d.save(p);return p
 def _claim_name(self,cid):
  with self._connect() as c:r=c.execute("SELECT condition_name FROM claims WHERE claim_id=?",(cid,)).fetchone();return r[0] if r else ""
 def _ensure(self):
  with self._connect() as c:c.executescript(NEXUS_SCHEMA)
 def _version(self,c,lid,num,content,meta):c.execute("INSERT INTO nexus_letter_versions VALUES(?,?,?,?,?,?)",(str(uuid.uuid4()),lid,num,json.dumps(content),json.dumps(meta),_now()))
 @staticmethod
 def _content(value):
  value=value or {};return {k:value.get(k,"" if k!="missing_facts" else []) for k in CONTENT_FIELDS}
 @staticmethod
 def _validate(status,primary,theories):
  if status not in NEXUS_STATUSES:raise ValueError("Unsupported nexus status")
  if primary not in NEXUS_THEORIES or any(x not in NEXUS_THEORIES for x in theories):raise ValueError("Unsupported nexus theory")
