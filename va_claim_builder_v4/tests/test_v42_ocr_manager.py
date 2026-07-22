from __future__ import annotations

from pathlib import Path

from core.documents import DocumentManager
from core.ocr import OCRManager
from core.projects import ProjectManager
from core.projects.paths import AppPaths


def _paths(tmp_path: Path) -> AppPaths:
    root = tmp_path / "app"
    return AppPaths(
        root=root,
        projects=root / "Projects",
        logs=root / "Logs",
        backups=root / "Backups",
        settings_file=root / "settings.json",
    ).ensure()


def test_plain_text_extraction_is_persistent_and_page_labeled(tmp_path: Path) -> None:
    project = ProjectManager(_paths(tmp_path)).create_project("OCR Test")
    source = tmp_path / "record.txt"
    source.write_text("Diagnosis: migraine\nTreatment: medication", encoding="utf-8")
    document, created = DocumentManager(project).import_file(source)
    assert created is True

    result = OCRManager(project).process(document.document_id)

    assert result.page_count == 1
    assert result.method == "plain-text"
    assert result.text_path.exists()
    text = result.text_path.read_text(encoding="utf-8")
    assert "--- PAGE 1 ---" in text
    assert "Diagnosis: migraine" in text

    reopened = OCRManager(project).get(document.document_id)
    assert reopened.character_count == len(text)
    assert DocumentManager(project).get(document.document_id).status == "ocr_complete"


def test_word_extraction_preserves_paragraph_text(tmp_path: Path) -> None:
    from docx import Document

    project = ProjectManager(_paths(tmp_path)).create_project("Word OCR Test")
    source = tmp_path / "statement.docx"
    word = Document()
    word.add_paragraph("First paragraph")
    word.add_paragraph("Second paragraph")
    word.save(source)
    document, _ = DocumentManager(project).import_file(source)

    result = OCRManager(project).process(document.document_id)
    text = result.text_path.read_text(encoding="utf-8")

    assert result.method == "docx-text"
    assert "First paragraph" in text
    assert "Second paragraph" in text
