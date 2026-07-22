from __future__ import annotations
from pathlib import Path
from typing import Any
from docx import Document
import json

class ClaimDraftingIntelligence:
    """Generates conservative review drafts from an approved fact matrix."""

    def build_sections(self, fact_matrix: dict[str, Any], document_type: str, starting_text: str = "") -> dict[str, Any]:
        condition = fact_matrix["condition_name"]
        facts = fact_matrix.get("facts_by_element", {})
        sections: list[dict[str, Any]] = []
        if starting_text.strip():
            sections.append({"heading": "Original Draft Starting Point", "paragraphs": [starting_text.strip()]})

        sections.append({"heading": "Verified and Approved Facts", "paragraphs": self._fact_paragraphs(facts) or ["No approved favorable facts are available yet."]})
        if fact_matrix.get("verified_timeline") or fact_matrix.get("reported_timeline"):
            sections.append({"heading": "Historical Timeline", "paragraphs": self._timeline_paragraphs(fact_matrix)})

        if document_type == "doctor_nexus":
            sections.extend(self._doctor_sections(condition, fact_matrix))
        elif document_type == "buddy_letter":
            sections.extend(self._buddy_sections(condition, fact_matrix))
        else:
            sections.extend(self._personal_sections(condition, fact_matrix))

        if fact_matrix.get("limiting_evidence"):
            sections.append({"heading": "Evidence Requiring Reconciliation", "paragraphs": [f"{x['proposition']} ({x['citation']})" for x in fact_matrix["limiting_evidence"]]})
        sections.append({"heading": "Required Review", "paragraphs": fact_matrix.get("guardrails", [])})
        return {"condition_name": condition, "document_type": document_type, "sections": sections, "source_annotation_ids": [f["annotation_id"] for rows in facts.values() for f in rows]}

    def export(self, package: dict[str, Any], output_dir: Path) -> list[Path]:
        output_dir = Path(output_dir); output_dir.mkdir(parents=True, exist_ok=True)
        safe = "".join(c if c.isalnum() else "_" for c in package["condition_name"]).strip("_")
        stem = f"{safe}_{package['document_type']}_Review_Draft"
        docx_path = output_dir / f"{stem}.docx"; json_path = output_dir / f"{stem}.json"
        doc = Document(); doc.add_heading(f"{package['condition_name']} — {package['document_type'].replace('_',' ').title()}", 0)
        doc.add_paragraph("Generated from approved evidence for human review. This is not a signed medical opinion or adopted witness statement.")
        for section in package["sections"]:
            doc.add_heading(section["heading"], 1)
            for paragraph in section["paragraphs"]:
                doc.add_paragraph(paragraph)
        doc.save(docx_path); json_path.write_text(json.dumps(package, indent=2, default=str), encoding="utf-8")
        return [docx_path, json_path]

    @staticmethod
    def _fact_paragraphs(facts: dict[str, list[dict[str, Any]]]) -> list[str]:
        out=[]
        for element, rows in facts.items():
            for f in rows:
                out.append(f"{element.replace('_',' ').title()}: {f['proposition']} ({f['citation']})")
        return out

    @staticmethod
    def _timeline_paragraphs(matrix: dict[str, Any]) -> list[str]:
        out=[]
        for label, rows in (("Verified", matrix.get("verified_timeline", [])), ("Reported", matrix.get("reported_timeline", []))):
            for e in rows:
                date=e.get("event_date_start") or "Date not established"
                out.append(f"{date} — {label}: {e.get('description','')}")
        return out

    @staticmethod
    def _doctor_sections(condition: str, matrix: dict[str, Any]) -> list[dict[str, Any]]:
        return [
            {"heading": "Question for the Reviewing Clinician", "paragraphs": [f"After reviewing the cited records and the veteran's reported history, please independently determine whether {condition} is related to service under the identified theory. Select and explain only the opinion you can medically support."]},
            {"heading": "Clinician Rationale Workspace", "paragraphs": ["Address relevant positive and negative evidence, competing causes, chronology, and why the medical principles apply to this veteran. Do not rely solely on the veteran's conclusion or this draft."]},
            {"heading": "Clinician Certification", "paragraphs": ["Records reviewed: ____________________", "Opinion selected: ____________________", "Rationale: ____________________", "Name, credentials, signature, and date: ____________________"]},
        ]

    @staticmethod
    def _buddy_sections(condition: str, matrix: dict[str, Any]) -> list[dict[str, Any]]:
        return [{"heading": "Witness Personalization Prompts", "paragraphs": [f"Describe when you first personally observed signs of {condition}.", "Describe specific episodes, frequency, and changes you personally witnessed.", "Explain effects on work, family, activities, sleep, or daily functioning.", "Do not diagnose the condition or offer a medical causation opinion.", "State how you know the veteran and the period covered by your observations."]}, {"heading": "Witness Certification", "paragraphs": ["I certify that this statement is true and based on my personal knowledge and observations.", "Name, relationship, signature, and date: ____________________"]}]

    @staticmethod
    def _personal_sections(condition: str, matrix: dict[str, Any]) -> list[dict[str, Any]]:
        return [{"heading": "Personal Narrative Framework", "paragraphs": [f"Explain the earliest symptoms or worsening of {condition} using the timeline above.", "Distinguish remembered onset from the date of formal diagnosis or treatment.", "Describe frequency, severity, treatment, self-treatment, and functional impact.", "Explain treatment gaps or delayed reporting truthfully where relevant.", "Use personal observations; leave medical causation conclusions to a qualified clinician."]}, {"heading": "Veteran Certification", "paragraphs": ["I certify that the foregoing is true and correct to the best of my knowledge and belief.", "Signature and date: ____________________"]}]
