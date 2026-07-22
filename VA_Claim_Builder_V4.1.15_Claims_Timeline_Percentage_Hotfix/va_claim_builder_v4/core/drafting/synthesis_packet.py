from __future__ import annotations
import json
from pathlib import Path
from docx import Document

class SynthesisPacketGenerator:
    def __init__(self, output_dir:Path): self.output_dir=Path(output_dir); self.output_dir.mkdir(parents=True,exist_ok=True)
    def generate(self, synthesis:dict):
        safe=''.join(c if c.isalnum() else '_' for c in synthesis['condition_name']).strip('_')
        docx_path=self.output_dir/f"{safe}_Evidence_Synthesis.docx"; json_path=self.output_dir/f"{safe}_Evidence_Synthesis.json"
        doc=Document(); doc.add_heading(f"Evidence Synthesis — {synthesis['condition_name']}",0)
        doc.add_paragraph(f"Theory: {synthesis.get('theory','unknown')}"); doc.add_paragraph(synthesis['guardrail'])
        doc.add_heading('Clear Evidence Ties',1)
        for item in synthesis['clear_ties']:
            doc.add_heading(item['claim_element'].replace('_',' ').title(),2); doc.add_paragraph(item['proposition']); doc.add_paragraph(f"Strength: {item['strength']} ({item['score']:.0%})"); doc.add_paragraph(item['why_it_matters']); doc.add_paragraph('Sources: '+', '.join(item['sources']))
        doc.add_heading('Timeline Reconciliation',1); t=synthesis['timeline']; doc.add_paragraph(f"Status: {t['status']} | Score: {t['score']:.0%} | Earliest supported onset: {t.get('earliest_supported_onset') or 'Not established'}")
        for b in synthesis['bridge_statements']: doc.add_paragraph(b,style='List Bullet')
        doc.add_heading('Missing Elements / Targeted Development',1)
        for m in synthesis['missing_elements'] or ['None identified from the configured theory.']: doc.add_paragraph(str(m).replace('_',' ').title(),style='List Bullet')
        doc.add_heading('Open Conflicts',1)
        for c in synthesis['open_conflicts'] or [{'summary':'None'}]: doc.add_paragraph(c.get('summary','None'),style='List Bullet')
        doc.add_heading('Negative or Limiting Context',1)
        for n in synthesis['negative_context'] or [{'proposition':'None identified in approved evidence.'}]: doc.add_paragraph(n.get('proposition',''),style='List Bullet')
        doc.save(docx_path); json_path.write_text(json.dumps(synthesis,indent=2,default=str),encoding='utf-8'); return [docx_path,json_path]
