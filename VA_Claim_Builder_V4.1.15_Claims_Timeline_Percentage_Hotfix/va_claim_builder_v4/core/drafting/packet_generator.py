from __future__ import annotations
import json
from datetime import datetime, timezone
from pathlib import Path
from docx import Document
from docx.shared import Inches

class PacketGenerator:
    """Generates review packets from approved evidence only."""
    def __init__(self, output_root: str | Path):
        self.output_root = Path(output_root)
        self.output_root.mkdir(parents=True, exist_ok=True)

    def generate_review_packet(self, claim_name: str, packet_type: str, revised_text: str,
                               approved_annotations: list[dict], timeline_events: list[dict],
                               unresolved_questions: list[str]) -> tuple[Path, dict]:
        safe = "".join(c if c.isalnum() or c in "-_ " else "_" for c in claim_name).strip().replace(" ", "_")
        path = self.output_root / f"{safe}_{packet_type}_review.docx"
        doc = Document()
        doc.add_heading(f"{claim_name} — {packet_type.replace('_',' ').title()} Review Packet", 0)
        doc.add_paragraph("DRAFT FOR HUMAN REVIEW. This document is not a signed medical opinion and must not be submitted as final evidence until reviewed and completed by the appropriate author.")
        doc.add_heading("Revised Draft", level=1)
        for para in revised_text.split("\n\n"):
            if para.strip(): doc.add_paragraph(para.strip())
        doc.add_heading("Historical Timeline", level=1)
        for event in timeline_events:
            date_label = event.get("event_date_start") or "Date unknown"
            doc.add_paragraph(f"{date_label} — {event.get('description','')}", style="List Bullet")
        doc.add_heading("Approved Evidence Reference", level=1)
        for ann in approved_annotations:
            source = ann.get("document_name") or ann.get("document_id", "source")
            page = ann.get("page") or "?"
            doc.add_paragraph(f"{ann.get('finding','')} ({source}, p. {page})", style="List Bullet")
        doc.add_heading("Questions Requiring Author Resolution", level=1)
        if unresolved_questions:
            for q in unresolved_questions: doc.add_paragraph(q, style="List Bullet")
        else:
            doc.add_paragraph("None identified by the current review pass.")
        doc.add_heading("Review and Signature", level=1)
        doc.add_paragraph("I reviewed this draft, corrected it as necessary, and confirm that the final statements reflect my independent knowledge and judgment.")
        doc.add_paragraph("Name/Credentials: __________________________________________")
        doc.add_paragraph("Signature: __________________________________  Date: __________________")
        doc.save(path)
        manifest = {"generated_at": datetime.now(timezone.utc).isoformat(), "claim": claim_name,
                    "packet_type": packet_type, "approved_annotation_ids": [a.get("id") for a in approved_annotations],
                    "timeline_event_ids": [e.get("id") for e in timeline_events], "unresolved_question_count": len(unresolved_questions)}
        path.with_suffix(".manifest.json").write_text(json.dumps(manifest, indent=2), encoding="utf-8")
        return path, manifest
