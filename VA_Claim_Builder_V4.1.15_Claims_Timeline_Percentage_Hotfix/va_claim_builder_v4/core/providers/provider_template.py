from __future__ import annotations
from pathlib import Path
from typing import Any
from docx import Document
from docx.shared import Inches

class ProviderTemplateGenerator:
    """Creates provider-ready drafts without representing them as already issued by the clinic."""
    def build_docx(self, provider: dict[str, Any], condition: str, body_sections: list[tuple[str,str]], output: Path,
                   branding_mode: str = "professional_template", approved_logo_path: str | None = None) -> Path:
        if branding_mode not in {"professional_template", "provider_supplied_branding"}:
            raise ValueError("Unsupported branding mode")
        if branding_mode == "provider_supplied_branding" and not approved_logo_path:
            raise ValueError("Provider-supplied or expressly authorized logo file is required")
        doc = Document()
        section = doc.sections[0]
        header = section.header
        p = header.paragraphs[0]
        if approved_logo_path:
            p.add_run().add_picture(approved_logo_path, width=Inches(1.25))
        p.add_run(provider.get("practice_name", "Provider Practice")).bold = True
        p.add_run("\n" + provider.get("provider_name", "Reviewing Clinician"))
        credentials = provider.get("credentials", "")
        if credentials: p.add_run(", " + credentials)
        contact = " | ".join(x for x in [provider.get("address"), provider.get("phone"), provider.get("website")] if x)
        if contact: p.add_run("\n" + contact)
        doc.add_heading(f"Draft Medical Opinion for Clinician Review — {condition}", 0)
        doc.add_paragraph("DRAFT — NOT ISSUED BY THE PROVIDER OR PRACTICE UNTIL INDEPENDENTLY REVIEWED, EDITED, AND SIGNED.")
        doc.add_paragraph("Practice and contact information should be verified by the clinician's office. Branding is provider-supplied or a neutral professional template; this document does not reproduce or impersonate official letterhead without authorization.")
        for heading, text in body_sections:
            doc.add_heading(heading, level=1)
            doc.add_paragraph(text)
        doc.add_heading("Clinician Attestation", 1)
        doc.add_paragraph("I independently reviewed the listed records, corrected the draft as necessary, and adopt only the statements and opinion I have personally verified.")
        doc.add_paragraph("Opinion: ______________________________________________")
        doc.add_paragraph("Rationale:\n\n\n")
        doc.add_paragraph("Signature: __________________  Date: __________  License/NPI (optional): __________")
        output = Path(output); output.parent.mkdir(parents=True, exist_ok=True); doc.save(output)
        return output
