from __future__ import annotations
from pathlib import Path
from docx import Document

class CPExamPreparationGenerator:
    def build(self, claim: dict, annotations: list[dict], timeline: list[dict], conflicts: list[dict]) -> dict:
        cid = claim.get("id")
        approved = [a for a in annotations if a.get("claim_id") == cid and a.get("review_status") == "approved"]
        return {
            "claim": claim,
            "confirmed_facts": [a for a in approved if a.get("polarity") == "favorable"],
            "negative_or_limiting_evidence": [a for a in approved if a.get("polarity") in {"unfavorable", "ambiguous"}],
            "timeline": [t for t in timeline if t.get("claim_id") == cid],
            "items_to_clarify": [c for c in conflicts if c.get("status", "open") == "open" and (not c.get("claim_id") or c.get("claim_id") == cid)],
            "instructions": [
                "Describe symptoms truthfully in your own words.",
                "Distinguish typical days from flare-ups and worst days.",
                "Explain frequency, duration, functional impact, treatment, and medication effects.",
                "Do not exaggerate, minimize, memorize scripted answers, or guess at facts you do not remember.",
            ],
        }
    def export(self, packet: dict, output_dir: Path) -> Path:
        output_dir.mkdir(parents=True, exist_ok=True)
        name = (packet["claim"].get("condition_name") or "claim").replace("/", "-")
        path = output_dir / f"{name}_C&P_Factual_Preparation.docx"
        doc = Document(); doc.add_heading(f"C&P Examination Factual Preparation — {name}", 0)
        doc.add_paragraph("For truthful factual preparation only. This document is not a script and should not be submitted as medical evidence unless independently appropriate.")
        for title, key in [("Confirmed facts", "confirmed_facts"), ("Negative or limiting evidence", "negative_or_limiting_evidence"), ("Timeline", "timeline"), ("Items to clarify", "items_to_clarify")]:
            doc.add_heading(title, level=1)
            rows = packet.get(key, [])
            if not rows: doc.add_paragraph("None identified.")
            for row in rows:
                text = row.get("finding") or row.get("summary") or row.get("event") or str(row)
                doc.add_paragraph(text, style="List Bullet")
        doc.add_heading("Preparation principles", level=1)
        for item in packet["instructions"]: doc.add_paragraph(item, style="List Bullet")
        doc.save(path); return path
